import re
import os
from elasticsearch import Elasticsearch
import time
from docx import Document
from mapping_creation import create_mapping
import networks
import netaddr
import sys
import logging


FORMAT_NAMES = ['ip_host', 'ip_dest', 'port_dest', 'declaration']

es = Elasticsearch([{'host': 'elasticsearch', 'port': 9200}])
VPN = []

with open('vpns', 'r') as fp:
    [VPN.append(x.strip('\n')) for x in fp.readlines()]

def ip_iterator(text):
    iterable = re.findall(r'\d+\.\d+\.\d+\.\d+/\d{2}|\d+\.\d+\.\d+\.\d+', text)
    if 'vpn' in text.lower():
        iterable = VPN
    return iterable




def parse_first_table(table):
    data_rows = []
    for row in table.rows[1:]:
        for srcip in ip_iterator(row.cells[0].text):
            for dstip in ip_iterator(row.cells[1].text):
                for dstport in re.findall(r'(\d+\s*[-]\s*\d+)|(\d+)', row.cells[2].text):
                    if dstport[1] != '':
                        data_rows.append([srcip, dstip,dstport[1],row.cells[3].text])
                    else:
                        data_rows.append([srcip, dstip, dstport[0], row.cells[3].text])
    return data_rows

def send_error(column_number, column_string, fileinfo):
    element = {
        'added_time': int(round(time.time() * 1000)),
        'ticket_type': fileinfo[0],
        'ticket_number': fileinfo[1],
        'column': column_number,
        'column_string': column_string,
    }
    es.index(index='errors', doc_type='table', body=element)

def scan_broken_table(table, fileinfo=('CRQ', '000000')):
    rows = table.rows[1:]
    for row in rows:
        if len(ip_iterator(row.cells[0].text)) == 0:
            send_error(0, row.cells[0].text,fileinfo)
            continue
        if len(ip_iterator(row.cells[1].text)) == 0:
            send_error(1, row.cells[1].text,fileinfo)
            continue
        if len(re.findall(r'(\d+\s*[-]\s*\d+)|(\d+)', row.cells[2].text)) == 0:
            send_error(2, row.cells[2].text, fileinfo)
            continue
    return True

def iterator(path):
    for dirpath, dirnames, files in os.walk(path):
        for file in files:
            if file.endswith(('.docx', '.DOCX')):
                yield '/'.join([dirpath,file])

def get_file_info(filename):
    last_folder = filename.split('/')[-2]
    ticket_type = last_folder[:3].encode('utf-8', 'surrogateescape').decode('utf8','surrogateescape')
    ticket_number = last_folder[3:].encode('utf-8', 'surrogateescape').decode('utf8','surrogateescape')
    doc_creation_date = os.path.getmtime(filename)
    return [ticket_type, ticket_number, doc_creation_date]

def format_table(table, filename):
    dict_table = [dict([tpl for tpl in zip(FORMAT_NAMES, row)]) for row in table]
    ticket_type, ticket_number, doc_creation_date = get_file_info(filename)
    [el.update([('doc_creation_date', int(round(doc_creation_date*1000))),('added_time', int(round(time.time()*1000)))]) for el in dict_table]
    for el in dict_table:
        src_ip_network = netaddr.IPNetwork(el['ip_host'])
        dst_ip_network = netaddr.IPNetwork(el['ip_dest'])
        src_ip, dst_ip = src_ip_network.ip, dst_ip_network.ip
        el['ip_host'], el['ip_dest'] = str(src_ip), str(dst_ip)
        el['host_network'], el['dst_network'] = str(src_ip_network), str(dst_ip_network)
        el['dst_network_description'] = networks.check(dst_ip_network) or 'UNKNOWN'
        el['ticket_type'], el['ticket_number'] = ticket_type, ticket_number
        port = el['port_dest']
        el['ip_port_triple'] = ':'.join(list(map(str, [src_ip, dst_ip, port])))
    return dict_table

def error_catching(func):
    def wrapper(filename):
        try:
            func(filename)
            logging.info('SENT \t' + filename)
        except Exception as e:
            logging.error('\t'.join([filename, str(e)]))
    return wrapper


@error_catching
def scan_doc(filename):
    doc = Document(filename)
    try:
        table = parse_first_table(doc.tables[1])
    except:
        raise Exception('parsing table_error')
    dict_to_send_to_el = format_table(table,filename)
    counter = 0
    for el in dict_to_send_to_el:
        try:
            es.index(index='tickets', doc_type='first_table', body=el)
            counter += 1
        except:
            try:
                el['ticket_type'], el['ticket_number'] = 'CRQ', '0000000'
                es.index(index='tickets', doc_type='first_table', body=el)
                counter += 1
            except:
                pass
    if counter == 0:
        scan_broken_table(doc.tables[1])
        raise Exception('Table is corrupted')
    scan_broken_table(doc.tables[1], (el['ticket_type'], el['ticket_number']))

if __name__ == '__main__':
    logging.basicConfig(level=logging.ERROR, format='%(asctime)s - %(levelname)s - %(message)s')
    path = '/usr/share/tickets'
    names_list = set()
    if '--scan_all' in sys.argv:
        create_mapping(es, 'tickets')
        create_mapping(es, 'errors')
        for x in iterator(path):
            scan_doc(x)
            names_list.add(x)
        while True:
            for x in iterator(path):
                if x not in names_list:
                    scan_doc(x)
                    names_list.add(x)
            time.sleep(600)
    else:
        start_time = time.time()
        while True:
            for x in iterator(path):
                if get_file_info(x)[2] > start_time and x not in names_list:
                    scan_doc(x)
                    names_list.add(x)
            time.sleep(600)